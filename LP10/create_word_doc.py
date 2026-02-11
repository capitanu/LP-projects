#!/usr/bin/env python3
"""Create Word document with search details and article references."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Cautare articole stiintifice Open Access", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Research topic
doc.add_heading("Tema de cercetare", level=2)
doc.add_paragraph(
    "Obezitatea infantila si biomarkerii inflamatori - rolul inflamatiei cronice "
    "de grad scazut in obezitatea pediatrica si markerii biologici asociati "
    "(chitotriosidaza, CRP, IL-6, adiponectina, disfunctia tesutului adipos)."
)

# Search objective
doc.add_heading("Obiectivul cercetarii", level=2)
doc.add_paragraph(
    "Identificarea si analiza articolelor stiintifice Open Access care investigheaza "
    "relatia dintre obezitatea infantila si markerii inflamatori, cu accent pe: "
    "activitatea chitotriosidazei ca biomarker al inflamatiei, markerii inflamatori "
    "cronici (hs-CRP, IL-6, IL-10, adiponectina), complicatiile cardiometabolice "
    "ale obezitatii pediatrice si disfunctia tesutului adipos."
)

# Search details
doc.add_heading("Detaliile cautarii", level=2)
p = doc.add_paragraph()
p.add_run("Platforma de cautare: ").bold = True
p.add_run("PubMed Central (PMC) / Europe PMC")

p2 = doc.add_paragraph()
p2.add_run("Termeni de cautare: ").bold = True
p2.add_run(
    "childhood obesity, inflammatory biomarkers, chitotriosidase, BMI, CRP, "
    "adipose tissue dysfunction, pediatric obesity, metabolic syndrome, "
    "open access, free full text"
)

p3 = doc.add_paragraph()
p3.add_run("Criterii de selectie: ").bold = True
p3.add_run(
    "Articole Open Access (acces liber), publicate in reviste indexate, "
    "relevante pentru tema obezitate infantila si inflamatie, "
    "publicate in perioada 2022-2024."
)

p4 = doc.add_paragraph()
p4.add_run("Numar articole selectate: ").bold = True
p4.add_run("4")

# Articles
doc.add_heading("Articolele selectate", level=2)

articles = [
    {
        "nr": 1,
        "title": "Evaluation of Circulating Chitotriosidase Activity in Children with Obesity",
        "authors": "Taranu I, Iancu M, Lazea C, Alkhzouz C, Racataianu N, Catana CS, Mirea AM, Miclea D, Bolboaca SD, Drugan C",
        "journal": "Journal of Clinical Medicine, 2022; 11(13):3634",
        "doi": "https://doi.org/10.3390/jcm11133634",
        "pmc": "PMC9267881",
        "link": "https://pmc.ncbi.nlm.nih.gov/articles/PMC9267881/",
    },
    {
        "nr": 2,
        "title": "Chronic Inflammatory Markers in Overweight and Obese Children: A Cross-sectional Analytical Study",
        "authors": "Gokulakrishnan R, Delhikumar CG, Senthilkumar GP, Sahoo J, Kumar RR",
        "journal": "Indian Journal of Endocrinology and Metabolism, 2024; 28(5):542-547",
        "doi": "https://doi.org/10.4103/ijem.ijem_353_23",
        "pmc": "PMC11642504",
        "link": "https://pmc.ncbi.nlm.nih.gov/articles/PMC11642504/",
    },
    {
        "nr": 3,
        "title": "Pediatric Obesity: Complications and Current Day Management",
        "authors": "Vajravelu ME, Tas E, Arslanian S",
        "journal": "Life (Basel), 2023; 13(7):1591",
        "doi": "https://doi.org/10.3390/life13071591",
        "pmc": "PMC10381624",
        "link": "https://pmc.ncbi.nlm.nih.gov/articles/PMC10381624/",
    },
    {
        "nr": 4,
        "title": "Obesity and Adipose Tissue Dysfunction: From Pediatrics to Adults",
        "authors": "Menendez A, Wanczyk H, Walker J, Zhou B, Santos M, Finck C",
        "journal": "Genes (Basel), 2022; 13(10):1866",
        "doi": "https://doi.org/10.3390/genes13101866",
        "pmc": "PMC9601855",
        "link": "https://pmc.ncbi.nlm.nih.gov/articles/PMC9601855/",
    },
]

for art in articles:
    doc.add_heading(f"Articolul {art['nr']}", level=3)

    p = doc.add_paragraph()
    p.add_run("Titlu: ").bold = True
    p.add_run(art["title"])

    p = doc.add_paragraph()
    p.add_run("Autori/Sursa: ").bold = True
    p.add_run(art["authors"])

    p = doc.add_paragraph()
    p.add_run("Revista: ").bold = True
    p.add_run(art["journal"])

    p = doc.add_paragraph()
    p.add_run("DOI / Identificator: ").bold = True
    p.add_run(art["doi"])

    p = doc.add_paragraph()
    p.add_run("Link PMC: ").bold = True
    p.add_run(art["link"])

    doc.add_paragraph()

output = "/Users/calin.capitanu/repos/github.com/proiect-andrada/LP10/Cautare_articole_BalmezAndrada.docx"
doc.save(output)
print(f"Created Word document: {output}")
